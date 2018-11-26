#!/usr/bin/env python3
from mymodules import files
import os
from docx import Document


location = "" #path to source folder
filetype = ''
aFiles = []
root = os.listdir(location)

for item in root:
    if os.path.isdir(os.path.join(location, item)): #if item is a direcotry
        aViewControllers = files.get_view_and_controller(os.path.join(location, item))
        aUtils = files.get_artifact_by_type("util", os.path.join(location, item))
        foldername = item
        filename = item + "_doc.docx"
        file = os.path.join(location, filename)
        document = Document()
        document.add_heading(item, 1)
        table = document.add_table(rows=0, cols=5)
        table.style = 'TableGrid'
        aCells = ['View path','View file name','Controller path','Controller file name','Controller name']
        row = table.add_row()
        for index, cell in enumerate(aCells):
            row.cells[index].text  = cell

        for item in aViewControllers:
            row_cells = table.add_row().cells
            relpath = item[0].split(foldername)[1]
            row_cells[0].text = relpath
            row_cells[1].text = item[1]
            relpath = item[2].split(foldername)[1]
            row_cells[2].text = relpath
            row_cells[3].text = item[3]
            row_cells[4].text = item[4]

        document.add_page_break()
        document.add_heading("View Controllers", 1)
        aProcessedControllers = []
        for item in aViewControllers:
            if item[4] not in aProcessedControllers:
                aProcessedControllers.append(item[4])
                controller = os.path.join(item[2], item[3])
                fnnames = files.get_functions_from_js(controller)
                document.add_heading(item[4], 2)
                table = document.add_table(rows=1, cols=2)
                table.style = 'TableGrid'
                hdr_cells = table.rows[0].cells
                table.rows[0].style = "borderColor:red;background-color:gray"
                hdr_cells[0].text = 'Function'
                hdr_cells[1].text = 'Description'
                
                for function in fnnames:
                    print(function)
                    row_cells = table.add_row().cells
                    row_cells[0].text = function
                    row_cells[1].text = ""
                
        aControllers = files.get_artifact_by_type("controller.js", os.path.join(location, foldername))
        for item in aControllers:
            if item[3] not in aProcessedControllers:
                aProcessedControllers.append(item[3])
                controller = os.path.join(item[0], item[1])
                fnnames = files.get_functions_from_js(controller)
                document.add_heading(item[3], 2)
                table = document.add_table(rows=1, cols=2)
                table.style = 'TableGrid'
                hdr_cells = table.rows[0].cells
                table.rows[0].style = "borderColor:red;background-color:gray"
                hdr_cells[0].text = 'Function'
                hdr_cells[1].text = 'Description'
                
                for function in fnnames:
                    print(function)
                    row_cells = table.add_row().cells
                    row_cells[0].text = function
                    row_cells[1].text = ""
                
        document.add_heading("Utilities", 1)
        document.add_heading("Util Files", 2)
        table = document.add_table(rows=1, cols=2)
        table.style = 'TableGrid'
        hdr_cells = table.rows[0].cells
        table.rows[0].style = "borderColor:red;background-color:gray"
        hdr_cells[0].text = 'Location'
        hdr_cells[1].text = 'File'
        
        for item in aUtils:
            row_cells = table.add_row().cells
            relpath = item[0].split(foldername)[1]
            row_cells[0].text = relpath
            row_cells[1].text = item[1]
        
        aProcessedUtils = []
        for item in aUtils:
            if item[1] not in aProcessedUtils:
                aProcessedUtils.append(item[1])
                document.add_heading(item[1], 2)
                controller = os.path.join(item[2])
                fnnames = files.get_functions_from_js(controller)
                table = document.add_table(rows=1, cols=2)
                table.style = 'TableGrid'
                hdr_cells = table.rows[0].cells
                table.rows[0].style = "borderColor:red;background-color:gray"
                hdr_cells[0].text = 'Function'
                hdr_cells[1].text = 'Description'
        
                for function in fnnames:
                    row_cells = table.add_row().cells
                    row_cells[0].text = function
                    row_cells[1].text = ""
        
        document.save(file)
        # print(file)
